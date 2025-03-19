from curses.ascii import alt
import streamlit as st
from dotenv import load_dotenv

# Set page configuration
st.set_page_config(
    page_title="DREAM JOB CATHCER",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

import tempfile
import openai
from pandas import options
import os
import re
import pandas as pd
from docx import Document
import pdfplumber
import plotly.express as px
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException
from webdriver_manager.chrome import ChromeDriverManager
import time
import json
import random
from pathlib import Path  # Fix for "Path is not defined"
import logging  # Add logging module for logger
from datetime import datetime  # Add datetime module for datetime
try:
    import fitz  # PyMuPDF
except ImportError:
    print("Error: PyMuPDF not installed. Please install with: pip install pymupdf")
    exit(1)
try:
    import docx  # python-docx for .docx files
except ImportError:
    print("Warning: python-docx not installed. Please install with: pip install python-docx")
    docx = None
# Optional: tqdm for progress bars
try:
    from tqdm import tqdm
except ImportError:
    def tqdm(iterable, *args, **kwargs):
        return iterable


# Configure logger
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("forte_extractor.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("forte_extractor")  # Initialize logger
load_dotenv()


class ForteDocumentExtractor:
    """
    A comprehensive data extractor for Forte documents that extracts metadata, 
    content structure, tables, form fields, and text with spatial information.
    """
    def __init__(self, config=None):
        """
        Initialize the Forte Document Extractor.
        
        Args:
            config (dict, optional): Configuration options for extraction.
        """
        self.config = config or {}
        self.default_output_dir = "extracted_data"
        
        # Define regex patterns for various document entities.
        self.patterns = {
            "amount": r"\$\s*[\d,]+\.\d{2}",
            "date": r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b",
            "email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
            "phone": r"\b(?:\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b",
            "account_number": r"\bAcct(?:ount)?\s*(?:#|No|Number)?[:.\s]*(\d+[-\s]?\d+[-\s]?\d+[-\s]?\d+)\b",
            "invoice_number": r"\bInv(?:oice)?\s*(?:#|No|Number)?[:.\s]*([A-Z0-9]+(?:[-\s]?[A-Z0-9]+)*)\b"
        }
        
        # Table extraction parameters.
        self.table_params = self.config.get("table_params", {
            "vertical_strategy": "lines",
            "horizontal_strategy": "lines",
            "explicit_vertical_lines": [],
            "explicit_horizontal_lines": [],
            "snap_tolerance": 3,
            "join_tolerance": 3,
            "edge_min_length": 3,
            "min_words_vertical": 3,
            "min_words_horizontal": 1
        })
        
        logger.info("ForteDocumentExtractor initialized")
    
    def extract_document(self, filepath):
        """
        Extract all data from a Forte document.
        
        Args:
            filepath (str): Path to the Forte document.
        
        Returns:
            dict: Dictionary containing extracted data.
        """
        logger.info(f"Processing document: {filepath}")
        if not os.path.exists(filepath):
            logger.error(f"File not found: {filepath}")
            return {"error": "File not found"}
        
        try:
            file_ext = Path(filepath).suffix.lower()
            if file_ext in ['.pdf']:
                result = self._extract_pdf(filepath)
            elif file_ext in ['.txt']:
                result = self._extract_txt(filepath)
            elif file_ext in ['.docx']:
                result = self._extract_docx(filepath)
            elif file_ext in ['.doc']:
                result = self._extract_doc(filepath)
            else:
                logger.error(f"Unsupported file format: {file_ext}")
                return {"error": f"Unsupported file format: {file_ext}"}
            
            # Ensure the result has a "content" key.
            if result is None or "content" not in result:
                logger.error("Extraction failed: No content found.")
                return {"error": "Extraction failed: No content found."}
            
            logger.info(f"Document processed successfully: {len(result['content'])} pages/sections")
            return result
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=True)
            return {"error": str(e)}
    
    def _extract_pdf(self, filepath):
        """
        Extract data from a PDF file.
        
        Args:
            filepath (str): Path to the PDF file.
        
        Returns:
            dict: Extracted data.
        """
        try:
            doc = fitz.open(filepath)
            result = {
                "metadata": self._extract_metadata(doc),
                "content": [],
                "tables": [],
                "form_fields": [],
                "entities": {k: [] for k in self.patterns.keys()},
                "file_info": {
                    "filename": os.path.basename(filepath),
                    "filepath": filepath,
                    "filesize": os.path.getsize(filepath),
                    "file_type": "PDF",
                    "extraction_time": datetime.now().isoformat()
                },
                "summary": {}
            }
            
            for page_num, page in enumerate(tqdm(doc, desc="Processing pages")):
                page_data = self._process_page(page, page_num)
                result["content"].append(page_data["content"])
                # Collect tables.
                for table in page_data["tables"]:
                    table["page_number"] = page_num + 1
                    result["tables"].append(table)
                # Collect form fields.
                for field in page_data["form_fields"]:
                    field["page_number"] = page_num + 1
                    result["form_fields"].append(field)
                # Collect entities.
                for entity_type, entities in page_data["entities"].items():
                    for entity in entities:
                        entity["page_number"] = page_num + 1
                        result["entities"][entity_type].append(entity)
            
            result["structure"] = self._extract_document_structure(doc)
            result["summary"] = self._generate_summary(result)
            doc.close()
            return result
        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}", exc_info=True)
            return {"error": str(e), "content": []}
    
    def _extract_txt(self, filepath):
        """
        Extract data from a text file.
        
        Args:
            filepath (str): Path to the text file.
        
        Returns:
            dict: Extracted data.
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
            result = {
                "metadata": {
                    "title": os.path.basename(filepath),
                    "author": "Unknown",
                    "creation_date": datetime.fromtimestamp(os.path.getctime(filepath)).isoformat(),
                    "modification_date": datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat(),
                },
                "content": [{
                    "page_number": 1,
                    "text": text,
                    "blocks": [{
                        "type": "text",
                        "text": text,
                        "bbox": [0, 0, 1, 1],
                        "block_number": 1
                    }]
                }],
                "tables": [],
                "form_fields": [],
                "entities": {k: [] for k in self.patterns.keys()},
                "file_info": {
                    "filename": os.path.basename(filepath),
                    "filepath": filepath,
                    "filesize": os.path.getsize(filepath),
                    "file_type": "TXT",
                    "extraction_time": datetime.now().isoformat()
                },
                "summary": {}
            }
            result["entities"] = self._extract_entities_from_text(text)
            result["summary"] = self._generate_summary(result)
            return result
        except Exception as e:
            logger.error(f"Error extracting TXT: {str(e)}", exc_info=True)
            return {"error": str(e)}
    
    def _extract_docx(self, filepath):
        """
        Extract data from a DOCX file.
        
        Args:
            filepath (str): Path to the DOCX file.
        
        Returns:
            dict: Extracted data.
        """
        if docx is None:
            logger.error("python-docx not installed. Cannot process .docx files.")
            return {"error": "python-docx not installed. Cannot process .docx files."}
        try:
            document = docx.Document(filepath)
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            result = {
                "metadata": {
                    "title": os.path.basename(filepath),
                    "author": "Unknown",
                    "creation_date": datetime.fromtimestamp(os.path.getctime(filepath)).isoformat(),
                    "modification_date": datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat(),
                },
                "content": [{
                    "page_number": 1,
                    "text": full_text,
                    "blocks": [{"type": "text", "text": p, "block_number": i+1} for i, p in enumerate(paragraphs)]
                }],
                "tables": [],
                "form_fields": [],
                "entities": {k: [] for k in self.patterns.keys()},
                "file_info": {
                    "filename": os.path.basename(filepath),
                    "filepath": filepath,
                    "filesize": os.path.getsize(filepath),
                    "file_type": "DOCX",
                    "extraction_time": datetime.now().isoformat()
                },
                "summary": {}
            }
            # Extract tables from DOCX
            for i, table in enumerate(document.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                result["tables"].append({
                    "table_number": i + 1,
                    "page_number": 1,
                    "data": table_data,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0
                })
            result["entities"] = self._extract_entities_from_text(full_text)
            result["summary"] = self._generate_summary(result)
            return result
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}", exc_info=True)
            return {"error": str(e)}
    
    def _extract_doc(self, filepath):
        """
        Extract data from a DOC file.
        
        Args:
            filepath (str): Path to the DOC file.
        
        Returns:
            dict: Extracted data.
        """
        try:
            with open(filepath, 'rb') as file:
                content = file.read().decode('utf-8', errors='ignore')
            result = {
                "metadata": {
                    "title": os.path.basename(filepath),
                    "author": "Unknown",
                    "creation_date": datetime.fromtimestamp(os.path.getctime(filepath)).isoformat(),
                    "modification_date": datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat(),
                },
                "content": [{
                    "page_number": 1,
                    "text": content,
                    "blocks": [{
                        "type": "text",
                        "text": content,
                        "block_number": 1
                    }]
                }],
                "tables": [],
                "form_fields": [],
                "entities": self._extract_entities_from_text(content),
                "file_info": {
                    "filename": os.path.basename(filepath),
                    "filepath": filepath,
                    "filesize": os.path.getsize(filepath),
                    "file_type": "DOC",
                    "extraction_time": datetime.now().isoformat()
                },
                "summary": {}
            }
            result["summary"] = self._generate_summary(result)
            return result
        except Exception as e:
            logger.error(f"Error extracting DOC: {str(e)}", exc_info=True)
            return {"error": str(e)}
    
    def _extract_metadata(self, doc):
        """
        Extract metadata from a document object (e.g., a PDF).
        
        Args:
            doc: Document object.
        
        Returns:
            dict: Metadata dictionary.
        """
        metadata = {}
        try:
            if hasattr(doc, 'metadata'):
                pdf_metadata = doc.metadata
                metadata = {
                    "title": pdf_metadata.get("title", ""),
                    "author": pdf_metadata.get("author", ""),
                    "subject": pdf_metadata.get("subject", ""),
                    "creator": pdf_metadata.get("creator", ""),
                    "producer": pdf_metadata.get("producer", ""),
                    "creation_date": pdf_metadata.get("creationDate", ""),
                    "modification_date": pdf_metadata.get("modDate", ""),
                    "keywords": pdf_metadata.get("keywords", "")
                }
        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")
            metadata = {"title": "", "author": "", "creation_date": "", "modification_date": "", "error": str(e)}
        return metadata
    
    def _extract_document_structure(self, doc):
        """
        Extract document structure (e.g., table of contents, bookmarks).
        
        Args:
            doc: Document object.
        
        Returns:
            dict: Document structure information.
        """
        structure = {"toc": [], "bookmarks": []}
        try:
            toc = doc.get_toc()
            structure["toc"] = [{"level": level, "title": title, "page": page} for level, title, page in toc]
            # Extract bookmarks if available
            outline = doc.outline
            if outline:
                def process_outline(items, level=0):
                    result = []
                    for item in items:
                        if item.dest:
                            page = doc.resolve_link(item.dest)
                            page_number = page[0] + 1 if isinstance(page, tuple) else None
                        else:
                            page_number = None
                        bookmark = {"title": item.title, "level": level, "page": page_number}
                        if item.down:
                            bookmark["children"] = process_outline(item.down, level + 1)
                        result.append(bookmark)
                    return result
                structure["bookmarks"] = process_outline(outline)
        except Exception as e:
            logger.warning(f"Error extracting document structure: {str(e)}")
        return structure
    
    def _extract_entities_from_text(self, text):
        """
        Extract entities from text using defined regex patterns.
        
        Args:
            text (str): Text to analyze.
        
        Returns:
            dict: Dictionary with keys as entity types and values as lists of entities.
        """
        entities = {k: [] for k in self.patterns.keys()}
        for entity_type, pattern in self.patterns.items():
            for match in re.finditer(pattern, text):
                entities[entity_type].append({
                    "value": match.group(0),
                    "start": match.start(),
                    "end": match.end(),
                    "context": text[max(0, match.start()-50):min(len(text), match.end()+50)]
                })
        return entities
    
    def _generate_summary(self, result):
        """
        Generate a summary from the extracted data.
        
        Args:
            result (dict): Extracted data.
        
        Returns:
            dict: Summary information.
        """
        summary = {
            "page_count": len(result["content"]),
            "word_count": sum(len(page["text"].split()) for page in result["content"] if "text" in page),
            "character_count": sum(len(page["text"]) for page in result["content"] if "text" in page),
            "table_count": len(result["tables"]),
            "form_field_count": len(result["form_fields"]),
            "entity_counts": {etype: len(entities) for etype, entities in result["entities"].items()},
            "file_info": result.get("file_info", {}),
            "metadata": result.get("metadata", {})
        }
        return summary
    
    def save_to_json(self, data, output_file=None):
        """
        Save the extracted data to a JSON file.
        
        Args:
            data (dict): Data to save.
            output_file (str, optional): Path to the output file.
        
        Returns:
            str: Path to the saved JSON file.
        """
        if output_file is None:
            os.makedirs(self.default_output_dir, exist_ok=True)
            if "file_info" in data and "filename" in data["file_info"]:
                base_name = os.path.splitext(data["file_info"]["filename"])[0]
            else:
                base_name = f"extracted_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            output_file = os.path.join(self.default_output_dir, f"{base_name}_extracted.json")
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            logger.info(f"Extracted data saved to {output_file}")
            return output_file
        except Exception as e:
            logger.error(f"Error saving to JSON: {str(e)}", exc_info=True)
            return None
    
    def save_to_csv(self, data, output_dir=None):
        """
        Save extracted data (tables, entities, and text content) into CSV files.
        
        Args:
            data (dict): Extracted data.
            output_dir (str, optional): Directory where CSV files will be saved.
        
        Returns:
            dict: Dictionary with keys corresponding to each saved CSV file.
        """
        if output_dir is None:
            output_dir = self.default_output_dir
        os.makedirs(output_dir, exist_ok=True)
        if "file_info" in data and "filename" in data["file_info"]:
            base_name = os.path.splitext(data["file_info"]["filename"])[0]
        else:
            base_name = f"extracted_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        saved_files = {}
        try:
            # Save tables.
            if data["tables"]:
                tables_dir = os.path.join(output_dir, "tables")
                os.makedirs(tables_dir, exist_ok=True)
                for i, table in enumerate(data["tables"]):
                    table_file = os.path.join(tables_dir, f"{base_name}_table_{i+1}.csv")
                    df = pd.DataFrame(table["data"])
                    df.to_csv(table_file, index=False, header=False)
                    saved_files[f"table_{i+1}"] = table_file
            # Save entities.
            entities_data = []
            for entity_type, entities in data["entities"].items():
                for entity in entities:
                    entities_data.append({
                        "entity_type": entity_type,
                        "value": entity["value"],
                        "page_number": entity.get("page_number", 1),
                        "context": entity["context"]
                    })
            if entities_data:
                entities_file = os.path.join(output_dir, f"{base_name}_entities.csv")
                pd.DataFrame(entities_data).to_csv(entities_file, index=False)
                saved_files["entities"] = entities_file
            # Save text content.
            text_file = os.path.join(output_dir, f"{base_name}_text.txt")
            with open(text_file, 'w', encoding='utf-8') as f:
                for page in data["content"]:
                    f.write(f"--- Page {page.get('page_number', '?')} ---\n\n")
                    f.write(page.get("text", ""))
                    f.write("\n\n")
            saved_files["text"] = text_file
            
            logger.info(f"Extracted data saved to CSV files in {output_dir}")
            return saved_files
        except Exception as e:
            logger.error(f"Error saving to CSV: {str(e)}", exc_info=True)
            return saved_files

# Function to initialize WebDriver
def init_driver():
    options = webdriver.ChromeOptions()
    options.add_argument("--headless")  # Run in background
    options.add_argument("--start-maximized")
    options.add_argument("--disable-blink-features=AutomationControlled")  # Bypass detection
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64)")
    options.add_argument("--disable-dev-shm-usage")  # Overcome limited resource problems
    options.add_argument("--no-sandbox")  # Bypass OS security model
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=options)

# Function to find element with retry for stale element issues
def find_element_safely(driver, by, value, max_attempts=3, wait_time=10):
    for attempt in range(max_attempts):
        try:
            # Wait for element to be present
            element = WebDriverWait(driver, wait_time).until(
                EC.presence_of_element_located((by, value))
            )
            return element
        except StaleElementReferenceException:
            if attempt == max_attempts - 1:
                raise
            print(f"Stale element, retrying... ({attempt+1}/{max_attempts})")
            time.sleep(1)
    return None  # Return None if element is not found

# Define the find_elements_safely function
def find_elements_safely(driver, by, value, max_attempts=3, wait_time=10):
    for attempt in range(max_attempts):
        try:
            # Wait for elements to be present
            elements = WebDriverWait(driver, wait_time).until(
                EC.presence_of_all_elements_located((by, value))
            )
            return elements
        except StaleElementReferenceException:
            if attempt == max_attempts - 1:
                raise
            print(f"Stale elements, retrying... ({attempt+1}/{max_attempts})")
            time.sleep(1)
    return []

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = "\n".join(paragraphs_text + tables_text)
        return all_text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

# CV Parsing Functions
def extract_name(text):
    lines = text.split("\n")
    possible_names = [line for line in lines[:5] if len(line.split()) in [2, 3]]  # Check first 5 lines
    return possible_names[0] if possible_names else "Unknown"

def extract_email(text):
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else "Not Found"

def extract_phone(text):
    match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    return match.group(0) if match else "Not Found"

def extract_location(text):
    # Look for location information in the first part of the resume
    lines = text.split("\n")
    location_section_start = 0
    for i, line in enumerate(lines):
        if "Details" in line:
            location_section_start = i
            break
    
    # Check a few lines after "Details" for location information
    for i in range(location_section_start, min(location_section_start + 10, len(lines))):
        if i < len(lines) and lines[i] in ["United States", "USA", "U.S.", "U.S.A."]:
            city_line = i - 1 if i > 0 else 0
            return f"{lines[city_line]}, {lines[i]}"
    
    # If not found with the above method, try pattern matching
    city_state_pattern = re.search(r"([A-Za-z\s]+),?\s+([A-Za-z\s]+)", text[:500])
    if city_state_pattern:
        return city_state_pattern.group(0)
    
    return "Location Not Found"

def extract_skills(text):
    # Expanded skill keywords based on the CV
    skill_keywords = [
        "Python", "Flask", "Django", "Machine Learning", "SQL", "Java", "React", "AWS",
        "JavaScript", "HTML", "CSS", "TensorFlow", "Pandas", "NumPy", "Docker", "Kubernetes",
        "Git", "Azure", "Linux", "Node.js", "C#", "C++", "Go", "PHP", "TypeScript", "Tableau", 
        "Power BI", "Jupyter", "Spark", "Hadoop", "Scala", "Cloud", "Vagrant", "LLMs", "GPT",
        "Re-enforcement", "Site Reliability", "DevOps", "Microservices", "NOSQL", 
        "Apache Kafka", "Apache Webserver", "Blockchain", "Performance engineering",
        "AI model Training", "Data Science", "Feature Engineering", "AI", "Shell Script", 
        "Intrusion Detection", "Matlab", "R", "Agile", "SDLC"
    ]
    
    # Extract skills section
    skills_section_match = re.search(r"Skills\n(.*?)(?:\n\n|\nProfile|\nEmployment)", text, re.DOTALL | re.IGNORECASE)
    skills_text = skills_section_match.group(1) if skills_section_match else text
    
    # Find skills from the predefined list
    found_skills = [skill for skill in skill_keywords if re.search(rf"\b{skill}\b", text, re.IGNORECASE)]
    
    # Also add skills that might be listed in the skills section but not in our predefined list
    if skills_section_match:
        additional_skills = [line.strip() for line in skills_text.split("\n") if line.strip() and len(line.strip()) > 2]
        # Combine and remove duplicates
        all_skills = list(set(found_skills + additional_skills))
    else:
        all_skills = found_skills
    
    return all_skills if all_skills else ["Skills not found"]

def extract_languages(text):
    # Look for a languages section
    languages_section = re.search(r"Languages\n(.*?)(?:\n\n|\nEducation|\nEmployment)", text, re.DOTALL | re.IGNORECASE)
    
    if languages_section:
        # Extract languages from the section
        languages_text = languages_section.group(1)
        languages = [lang.strip() for lang in languages_text.split("\n") if lang.strip()]
        return languages
    
    # Fallback to common language names if section not found
    common_languages = ["English", "Spanish", "French", "German", "Chinese", "Japanese", 
                        "Arabic", "Hindi", "Bengali", "Russian", "Portuguese"]
    found_languages = [lang for lang in common_languages if re.search(rf"\b{lang}\b", text)]
    
    return found_languages if found_languages else ["Not Found"]

def extract_experience(text):
    # Extracting years of experience patterns
    years_of_experience = re.findall(r"(\d{1,2})\+?\s*years?(?:\s+of\s+experience)?", text)
    
    if years_of_experience:
        max_experience = max(map(int, years_of_experience))
        return f"{max_experience} years of experience"
    
    # Alternative approach: count employment duration
    employment_pattern = re.findall(r"(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", text, re.IGNORECASE)
    
    if employment_pattern:
        return f"{len(employment_pattern)} employment periods found"
        
    return "Experience not found"

def extract_qualifications(text):
    # Expanded list of qualification keywords
    qualification_keywords = [
        "Bachelor's", "Master's", "PhD", "Degree", "Certification", "Diploma", "BSc", "MSc", 
        "B.A.", "M.A.", "BTech", "MTech", "MBA", "Engineering", "Architecture", "Computer Science", 
        "Information Technology", "Data Science", "Machine Learning", "AI", "Software Engineering",
        "CSSA", "Certified", "Business", "MIS", "University"
    ]
    
    # Extract education section
    education_section = re.search(r"Education\n(.*?)(?:\n\n|$)", text, re.DOTALL | re.IGNORECASE)
    
    if education_section:
        edu_text = education_section.group(1)
        # Extract full education entries
        education_entries = [line.strip() for line in edu_text.split("\n") if line.strip() 
                            and any(keyword in line for keyword in ["Bachelor", "Master", "PhD", "Degree"])]
        
        if education_entries:
            return education_entries
    
    # Fallback to keyword matching
    found_qualifications = []
    for qual in qualification_keywords:
        if re.search(rf"\b{qual}\b", text, re.IGNORECASE):
            # Get the context around the qualification keyword
            context_match = re.search(rf".{{0,50}}\b{qual}\b.{{0,50}}", text, re.IGNORECASE)
            if context_match:
                found_qualifications.append(context_match.group(0).strip())
            else:
                found_qualifications.append(qual)
    
    return list(set(found_qualifications)) if found_qualifications else ["Not Found"]

def extract_employment_history(text):
    # Find employment history section
    employment_section = re.search(r"Employment History\n(.*?)(?:\nEducation|\n\nEducation|$)", text, re.DOTALL | re.IGNORECASE)
    
    if not employment_section:
        # Try alternative headings
        employment_section = re.search(r"(?:Work Experience|Professional Experience|Career History)\n(.*?)(?:\nEducation|\n\nEducation|$)", text, re.DOTALL | re.IGNORECASE)
        
    if not employment_section:
        return ["Employment history not found"]
    
    employment_text = employment_section.group(1)
    
    # Extract job entries based on date patterns
    # Pattern looks for: Job Title, Company, Location then a date range on the next line
    job_pattern = re.findall(r"(.*?),\s+(.*?),\s+(.*?)\n(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", 
                             employment_text, re.IGNORECASE | re.MULTILINE)
    
    # If the standard pattern fails, try an alternative pattern
    if not job_pattern:
        job_pattern = re.findall(r"(.*?),\s+(.*?)\n(\w+\s+\d{4})\s*[—–-]\s*(\w+\s+\d{4}|PRESENT)", 
                                employment_text, re.IGNORECASE | re.MULTILINE)
    
    # If still no match, just try to get job titles
    if not job_pattern:
        # Look for lines that might be job titles followed by dates
        lines = employment_text.split("\n")
        jobs = []
        for i, line in enumerate(lines):
            if re.search(r"\b(architect|analyst|lead|principal|manager|director|engineer)\b", line, re.IGNORECASE):
                if i < len(lines) - 1 and re.search(r"\d{4}", lines[i+1]):
                    jobs.append(line)
        
        return jobs if jobs else ["Could not parse employment details"]
    
    # Format the job entries
    formatted_jobs = []
    for job in job_pattern:
        if len(job) == 5:  # Full pattern match
            formatted_jobs.append(f"{job[0]} at {job[1]}, {job[2]} ({job[3]} - {job[4]})")
        elif len(job) == 4:  # Alternative pattern match
            formatted_jobs.append(f"{job[0]} at {job[1]} ({job[2]} - {job[3]})")
    
    return formatted_jobs

def extract_profile_summary(text):
    # Find profile section
    profile_section = re.search(r"Profile\n(.*?)(?:\nEmployment|\n\nEmployment|\nSkills|\n\nSkills)", text, re.DOTALL | re.IGNORECASE)
    
    if profile_section:
        profile_text = profile_section.group(1).strip()
        # Trim to a reasonable length if too long
        if len(profile_text) > 500:
            return profile_text[:497] + "..."
        return profile_text
    
    # Alternative: look for summary or about section
    alt_section = re.search(r"(?:Summary|About|Personal Statement|Objective)\n(.*?)(?:\n\n|\nSkills|\nExperience)", text, re.DOTALL | re.IGNORECASE)
    
    if alt_section:
        return alt_section.group(1).strip()
    
    # Last resort: take first paragraph if short enough
    paragraphs = text.split("\n\n")
    if paragraphs and len(paragraphs[0]) < 200:
        return paragraphs[0].strip()
    
    return "Profile summary not found"

def extract_certifications(text):
    # Look for certifications section
    cert_section = re.search(r"(?:Certifications|Certificates|Professional Development)\n(.*?)(?:\n\n|\nSkills|\nExperience|$)", text, re.DOTALL | re.IGNORECASE)
    
    if cert_section:
        cert_text = cert_section.group(1)
        certifications = [cert.strip() for cert in cert_text.split("\n") if cert.strip()]
        return certifications
    
    # Look for common certification keywords
    cert_keywords = ["Certified", "Certificate", "Certification", "CISSP", "PMP", "AWS Certified", 
                    "Microsoft Certified", "Google Certified", "CompTIA", "ITIL", "Six Sigma"]
    
    found_certs = []
    for keyword in cert_keywords:
        cert_matches = re.findall(rf"{keyword}[^.\n]*", text, re.IGNORECASE)
        found_certs.extend(cert_matches)
    
    return found_certs if found_certs else ["No certifications found"]

def extract_social_links(text):
    linkedin_match = re.search(r"(?:linkedin\.com|LinkedIn:)[^\s]*", text)
    github_match = re.search(r"(?:github\.com|GitHub:)[^\s]*", text)
    portfolio_match = re.search(r"(?:Portfolio:|Website:)[^\s]*", text)
    
    social_links = {}
    if linkedin_match:
        social_links["LinkedIn"] = linkedin_match.group(0)
    if github_match:
        social_links["GitHub"] = github_match.group(0)
    if portfolio_match:
        social_links["Portfolio"] = portfolio_match.group(0)
    
    return social_links if social_links else {"Links": "No social links found"}

def extract_projects(text):
    # Look for projects section
    projects_section = re.search(r"(?:Projects|Personal Projects|Key Projects)\n(.*?)(?:\n\n|\nEducation|\nEmployment|$)", text, re.DOTALL | re.IGNORECASE)
    
    if projects_section:
        projects_text = projects_section.group(1)
        # Split by lines and look for project entries
        lines = projects_text.split("\n")
        projects = []
        current_project = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # New project might start with capitalized words
            if re.match(r"^[A-Z][a-z]", line) and not current_project.endswith(",") and current_project:
                projects.append(current_project)
                current_project = line
            else:
                if current_project:
                    current_project += " " + line
                else:
                    current_project = line
        
        # Add the last project
        if current_project:
            projects.append(current_project)
            
        return projects
    
    # Alternative approach: look for project keywords in experience section
    project_keywords = ["Project:", "Developed", "Built", "Created", "Implemented", "Designed"]
    project_matches = []
    
    for keyword in project_keywords:
        matches = re.findall(rf"{keyword}[^.\n]*(?:[.\n]|$)", text, re.IGNORECASE)
        project_matches.extend(matches)
    
    return project_matches if project_matches else ["No projects found"]



# Function to extract job details from a specific job page
def get_job_details(driver, apply_link):
    try:
        # Navigate to job details page
        driver.get(apply_link)
        # Longer random wait to ensure page loads completely
        time.sleep(random.uniform(4, 7))
        
        job_details = {}
        job_details["Structured Content"] = extract_structured_job_content(driver)
        job_description = extract_job_description(driver)
        job_details["Job Description Details"] = {
            "html": job_description.get("html", ""),
            "text": job_description.get("text", ""),
        }
        job_details["Structured Content"] = extract_structured_job_content(driver)

        # Extract job title
        try:
            job_title = find_element_safely(driver, By.CSS_SELECTOR, "h1.chakra-heading[data-testid='viewJobTitle']", wait_time=15).text
            job_details["Job Title"] = job_title
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            print(f"Could not extract job title: {str(e)[:50]}")
            job_details["Job Title"] = "Not available"
        
        # Extract company name
        try:
            company_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span[data-testid='detailText']", wait_time=15)
            if company_elements:
                job_details["Company Name"] = company_elements[0].text
            else:
                job_details["Company Name"] = "Not available"
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            print(f"Could not extract company name: {str(e)[:50]}")
            job_details["Company Name"] = "Not available"
        
        # Extract job rating
        try:
            rating_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span[data-testid='viewJobCompanyRating']", wait_time=10)
            if rating_elements:
                rating_text = rating_elements[0].text.strip()
                if rating_text:
                    # Extract only the numeric part of the rating if it exists
                    import re
                    rating_match = re.search(r'(\d+\.\d+)', rating_text)
                    if rating_match:
                        job_details["Job Rating"] = rating_match.group(1)
                    else:
                        job_details["Job Rating"] = "Not specified"
                else:
                    job_details["Job Rating"] = "Not specified"
            else:
                job_details["Job Rating"] = "Not available"
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            print(f"Could not extract job rating: {str(e)[:50]}")
            job_details["Job Rating"] = "Not available"
        
        # Extract job location
        try:
            # Try multiple selectors for location
            location_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span.chakra-stack[data-testid='viewJobCompanyLocation'] span[data-testid='detailText']", wait_time=10)
            if not location_elements:
                location_elements = find_elements_safely(driver, By.CSS_SELECTOR, "div[data-testid='viewJobCompanyLocation'] span[data-testid='detailText']", wait_time=10)
            
            if location_elements:
                job_details["Job Location"] = location_elements[0].text
            else:
                # Try an alternative approach
                all_detail_texts = find_elements_safely(driver, By.CSS_SELECTOR, "span[data-testid='detailText']")
                for element in all_detail_texts:
                    text = element.text
                    # Look for texts that might be locations (containing state abbreviations or "remote")
                    if re.search(r'\b[A-Z]{2}\b', text) or "remote" in text.lower():
                        job_details["Job Location"] = text
                        break
                else:
                    job_details["Job Location"] = "Not available"
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            print(f"Could not extract job location: {str(e)[:50]}")
            job_details["Job Location"] = "Not available"
        
        # Extract salary information
        try:
            # Try multiple selectors for salary
            salary_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span.chakra-stack[data-testid='viewJobBodyJobCompensation'] span[data-testid='detailText']", wait_time=8)
            if not salary_elements:
                salary_elements = find_elements_safely(driver, By.CSS_SELECTOR, "div[data-testid='viewJobBodyJobCompensation'] span[data-testid='detailText']", wait_time=8)
            
            if salary_elements:
                job_details["Salary"] = salary_elements[0].text
            else:
                # Try to find any text containing salary indicators
                all_texts = driver.find_elements(By.XPATH, "//*[contains(text(), '$') or contains(text(), 'salary') or contains(text(), 'Salary') or contains(text(), 'compensation') or contains(text(), 'Compensation')]")
                for element in all_texts:
                    text = element.text
                    if '$' in text and len(text) < 100:  # Likely a salary not a long description
                        job_details["Salary"] = text
                        break
                else:
                    job_details["Salary"] = "Not disclosed"
        except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
            print(f"Could not extract salary: {str(e)[:50]}")
            job_details["Salary"] = "Not disclosed"
        
        # Extract job type
        try:
            job_type_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span[data-testid='detailText']")
            job_type_found = False
            for element in job_type_elements:
                try:
                    text = element.text
                    if any(job_type in text.lower() for job_type in ["full-time", "part-time", "contract", "temporary", "internship"]):
                        job_details["Job Type"] = text
                        job_type_found = True
                        break
                except StaleElementReferenceException:
                    continue
            
            if not job_type_found:
                # Try to find job type in the page content
                page_content = driver.find_element(By.TAG_NAME, "body").text.lower()
                if "full-time" in page_content:
                    job_details["Job Type"] = "Full-time"
                elif "part-time" in page_content:
                    job_details["Job Type"] = "Part-time"
                elif "contract" in page_content:
                    job_details["Job Type"] = "Contract"
                elif "temporary" in page_content:
                    job_details["Job Type"] = "Temporary"
                elif "internship" in page_content:
                    job_details["Job Type"] = "Internship"
                else:
                    job_details["Job Type"] = "Not specified"
        except (TimeoutException, NoSuchElementException) as e:
            print(f"Could not extract job type: {str(e)[:50]}")
            job_details["Job Type"] = "Not specified"
        
        # Extract posting time - this seems to be particularly problematic
        try:
            # Try multiple approaches to find posting time
            posting_time = None
            
            # First approach
            posting_time_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span.chakra-stack[data-testid='viewJobBodyJobPostingTimestamp'] span[data-testid='detailText']", wait_time=8)
            if posting_time_elements:
                posting_time = posting_time_elements[0].text
            
            # Second approach
            if not posting_time:
                posting_time_elements = find_elements_safely(driver, By.CSS_SELECTOR, "div[data-testid='viewJobBodyJobPostingTimestamp'] span[data-testid='detailText']", wait_time=8)
                if posting_time_elements:
                    posting_time = posting_time_elements[0].text
            
            # Third approach - look for text containing time indicators
            if not posting_time:
                # Look for text nodes that might contain posting time information
                possible_elements = driver.find_elements(By.XPATH, "//*[contains(text(), 'day') or contains(text(), 'week') or contains(text(), 'month') or contains(text(), 'hour') or contains(text(), 'minute') or contains(text(), 'Posted')]")
                for element in possible_elements:
                    text = element.text
                    if re.search(r'(posted|ago)', text.lower()) and len(text) < 100:
                        posting_time = text
                        break
            
            job_details["Posted"] = posting_time if posting_time else "Not available"
        except Exception as e:
            print(f"Could not extract posting time with all approaches: {str(e)[:50]}")
            job_details["Posted"] = "Not available"
        
        # Extract benefits - implement multiple fallback strategies
        try:
            benefits = []
            
            # First approach - try the specific test ID
            benefits_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span[data-testid='viewJobBenefitItem']", wait_time=8)
            if benefits_elements:
                for benefit in benefits_elements:
                    try:
                        benefits.append(benefit.text)
                    except StaleElementReferenceException:
                        continue
            
            # Second approach - look for content that might indicate benefits
            if not benefits:
                benefit_keywords = ["benefit", "insurance", "vacation", "pto", "401k", "health", "dental", "vision"]
                benefit_section = driver.find_elements(By.XPATH, "//div[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'benefit')]")
                
                if benefit_section:
                    # Try to extract benefits from surrounding elements
                    for section in benefit_section:
                        parent = section.find_element(By.XPATH, "./..")
                        items = parent.find_elements(By.TAG_NAME, "li")
                        for item in items:
                            benefits.append(item.text)
            
            job_details["Benefits"] = benefits
        except (TimeoutException, NoSuchElementException) as e:
            print(f"Could not extract benefits: {str(e)[:50]}")
            job_details["Benefits"] = []
        
        # Extract qualifications with multiple approaches
        try:
            qualifications = []
            
            # First approach - try the specific test ID
            qualification_elements = find_elements_safely(driver, By.CSS_SELECTOR, "span[data-testid='viewJobQualificationItem']", wait_time=8)
            if qualification_elements:
                for qual in qualification_elements:
                    try:
                        qualifications.append(qual.text)
                    except StaleElementReferenceException:
                        continue
            
            # Second approach - look for section that might contain qualifications
            if not qualifications:
                qual_keywords = ["qualification", "requirement", "skill", "experience", "education"]
                for keyword in qual_keywords:
                    qual_section = driver.find_elements(By.XPATH, f"//div[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{keyword}')]")
                    if qual_section:
                        for section in qual_section:
                            try:
                                parent = section.find_element(By.XPATH, "./..")
                                items = parent.find_elements(By.TAG_NAME, "li")
                                if items:
                                    for item in items:
                                        qualifications.append(item.text)
                                    break  # Found valid qualifications, no need to check further
                            except:
                                continue
                
            job_details["Qualifications"] = qualifications
        except (TimeoutException, NoSuchElementException) as e:
            print(f"Could not extract qualifications: {str(e)[:50]}")
            job_details["Qualifications"] = []
        
        # Extract full job description with multiple fallback strategies
        try:
            # Try multiple approaches to get the job description
            full_description = ""
            
            # First approach: using section title
            try:
                desc_section_elements = find_elements_safely(driver, By.CSS_SELECTOR, "h2.chakra-text[data-testid='viewJobDetailsSectionTitle']")
                if desc_section_elements:
                    desc_section = desc_section_elements[0]
                    # Get the parent container and then all text content
                    parent_element = desc_section.find_element(By.XPATH, "./..")
                    full_description = parent_element.text.replace("Full Job Description", "").strip()
            except (NoSuchElementException, StaleElementReferenceException) as e:
                print(f"Could not extract job description with first approach: {str(e)[:50]}")
            
            # Second approach: try to find description container directly
            if not full_description:
                try:
                    desc_container = find_element_safely(driver, By.CSS_SELECTOR, "div[data-testid='viewJobBodyJobFullDescriptionContent']")
                    full_description = desc_container.text.strip()
                except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
                    print(f"Could not extract job description with second approach: {str(e)[:50]}")
            
            # Third approach: try to find any main content section
            if not full_description:
                try:
                    # Look for a main content container
                    content_containers = [
                        "div.css-je9w1e",
                        "div[data-testid='viewJobBodyJobFullDescriptionContent']",
                        "div.job-description",
                        "div.description"
                    ]
                    
                    for selector in content_containers:
                        try:
                            elements = driver.find_elements(By.CSS_SELECTOR, selector)
                            if elements:
                                full_description = elements[0].text.strip()
                                break
                        except:
                            continue
                except Exception as e:
                    print(f"Could not extract job description with third approach: {str(e)[:50]}")
            
            # Fourth approach: if all else fails, try to get relevant sections from the page
            if not full_description:
                try:
                    # Get all paragraph elements that might contain job description
                    paragraphs = driver.find_elements(By.TAG_NAME, "p")
                    if paragraphs:
                        full_description = "\n\n".join([p.text for p in paragraphs if len(p.text) > 50])
                except Exception as e:
                    print(f"Could not extract job description with fourth approach: {str(e)[:50]}")
            
            job_details["Job Description"] = full_description if full_description else "Not available"
        except Exception as e:
            print(f"Failed to extract job description: {str(e)[:50]}")
            job_details["Job Description"] = "Not available"
        
        return job_details
    
    except Exception as e:
        print(f"Error extracting job details from {apply_link}: {str(e)[:50]}")
        return {"Error": str(e), "Apply Link": apply_link}

# Enhanced function to extract job description with better HTML parsing
def extract_job_description(driver):
    try:
        # First approach: Try to get the exact container with the job description
        description_container = find_element_safely(
            driver, 
            By.CSS_SELECTOR, 
            "div[data-testid='viewJobBodyJobFullDescriptionContent']", 
            wait_time=15
        )
        
        if description_container:
            # Extract the full HTML content to preserve structure
            description_html = description_container.get_attribute('innerHTML')
            
            # Also get plain text for backup
            description_text = description_container.text.strip()
            
            # Return both formats
            return {
                "html": description_html,
                "text": description_text
            }
        
        # Second approach: Try alternative selectors if the first one fails
        alternative_selectors = [
            "div.css-cxpe4v",
            "div.job-description",
            "div[data-testid*='description']",
            "div.css-je9w1e"
        ]
        
        for selector in alternative_selectors:
            try:
                element = find_element_safely(driver, By.CSS_SELECTOR, selector, wait_time=5)
                if element:
                    description_html = element.get_attribute('innerHTML')
                    description_text = element.text.strip()
                    return {
                        "html": description_html,
                        "text": description_text
                    }
            except:
                continue
        
        # Third approach: Try to extract structured content specifically
        description_sections = {}
        
        # Common section headers in job descriptions
        section_keywords = [
            "About", "Company", "Role", "Responsibilities", "Requirements", 
            "Qualifications", "Skills", "Benefits", "Offer", "Compensation"
        ]
        
        for keyword in section_keywords:
            # Look for headers containing these keywords
            header_elements = driver.find_elements(
                By.XPATH, 
                f"//b[contains(text(), '{keyword}')] | //strong[contains(text(), '{keyword}')] | //h3[contains(text(), '{keyword}')] | //h4[contains(text(), '{keyword}')]"
            )
            
            for header in header_elements:
                try:
                    # Get the parent node that might contain the whole section
                    parent = header.find_element(By.XPATH, "./ancestor::p | ./ancestor::div[position()=1]")
                    section_title = header.text.strip()
                    
                    # Try to get the next sibling elements until next header
                    section_content = []
                    
                    # If it's a list, try to get all list items
                    list_items = parent.find_elements(By.XPATH, "./following-sibling::ul[1]/li")
                    if list_items:
                        for item in list_items:
                            section_content.append("• " + item.text.strip())
                    else:
                        # Otherwise get the text content
                        section_content.append(parent.text.replace(section_title, "").strip())
                    
                    description_sections[section_title] = "\n".join(section_content)
                except:
                    continue
        
        # If we found structured sections, use them
        if description_sections:
            structured_text = "\n\n".join([f"{title}:\n{content}" for title, content in description_sections.items()])
            return {
                "html": "",
                "text": structured_text,
                "structured": description_sections
            }
        
        # Fourth approach: Last resort, look for any content
        all_paragraphs = driver.find_elements(By.TAG_NAME, "p")
        if all_paragraphs:
            paragraphs_text = "\n\n".join([p.text for p in all_paragraphs if len(p.text) > 20])
            return {
                "html": "",
                "text": paragraphs_text
            }
        
        return {
            "html": "",
            "text": "No job description found"
        }
    
    except Exception as e:
        print(f"Error extracting job description: {str(e)[:100]}")
        return {
            "html": "",
            "text": f"Error extracting job description: {str(e)[:100]}"
        }

# Function to extract all structured content from job listing
def extract_structured_job_content(driver):
    try:
        # Extract sections like Responsibilities, Requirements, Benefits
        structured_data = {}
        
        # Look for unordered lists which often contain requirements or responsibilities
        list_sections = driver.find_elements(By.CSS_SELECTOR, "ul")
        
        for list_section in list_sections:
            try:
                # Try to find a header before this list
                header = list_section.find_element(By.XPATH, "./preceding::b[1] | ./preceding::strong[1] | ./preceding::h3[1] | ./preceding::h4[1]")
                section_title = header.text.strip()
                
                # Get all list items
                items = list_section.find_elements(By.TAG_NAME, "li")
                section_items = [item.text.strip() for item in items if item.text.strip()]
                
                # Only add non-empty sections
                if section_title and section_items:
                    structured_data[section_title] = section_items
            except:
                # If no header is found, try to categorize by content
                items = list_section.find_elements(By.TAG_NAME, "li")
                section_items = [item.text.strip() for item in items if item.text.strip()]
                
                if section_items:
                    # Try to guess the section based on content
                    content = " ".join(section_items).lower()
                    if any(word in content for word in ["degree", "experience", "year", "skill", "proficiency"]):
                        structured_data["Requirements"] = section_items
                    elif any(word in content for word in ["responsible", "duties", "develop", "create", "manage"]):
                        structured_data["Responsibilities"] = section_items
                    elif any(word in content for word in ["benefit", "insurance", "vacation", "remote", "flexible"]):
                        structured_data["Benefits"] = section_items
        
        return structured_data
    
    except Exception as e:
        print(f"Error extracting structured content: {str(e)[:100]}")
        return {}


#Example workflow to extract skills and scrape jobs

# Initialize session state variables
if "uploaded_file" not in st.session_state:
    st.session_state.uploaded_file = None
if "uploaded_file_content" not in st.session_state:
    st.session_state.uploaded_file_content = None
if "cv_data" not in st.session_state:
    st.session_state.cv_data = None
if "job_results" not in st.session_state:
    st.session_state.job_results = None
if "scraping_in_progress" not in st.session_state:
    st.session_state.scraping_in_progress = False
if "process_completed" not in st.session_state:
    st.session_state.process_completed = False
if "current_skill" not in st.session_state:
    st.session_state.current_skill = None
if "current_page" not in st.session_state:
    st.session_state.current_page = 1

#Modified function to scrape jobs by skills
def scrape_jobs_by_skills(skills, driver):
    results = st.session_state.job_results or {}
    for skill in skills:
        if st.session_state.current_skill and st.session_state.current_skill != skill:
            continue  # Resume from the last skill being scraped
        st.session_state.current_skill = skill
        st.write(f"Scraping jobs for skill: **{skill}**")
        print(f"Scraping jobs for skill: {skill}")
        skill_results = scrape_single_job(skill, driver)
        results[skill] = skill_results
        st.session_state.current_skill = None  # Reset after completing the skill
    return results

# Modified function to scrape a single job role
def scrape_single_job(role, driver):
    results = []
    max_jobs = st.session_state.get("num_jobs_per_skill", 10)  # Get user-selected job limit
    max_pages = st.session_state.get("num_pages_to_scrape", 5)  # Get user-selected page limit
    job_count = 0
    current_page = 1
    try:
        search_url = f"https://www.simplyhired.com/search?q={role.replace(' ', '+')}&location=&countryCode=&page={st.session_state.current_page}&pageSize=1&language=en"
        while job_count < max_jobs and current_page <= max_pages:
            jobs = driver.find_elements(By.CLASS_NAME, "css-1djbb1k")
            time.sleep(random.uniform(5, 8))  # Prevent bot detection
            if not jobs:
                print(f"No more jobs found for {role}.")
            print(f"Processing page {st.session_state.current_page} for {role}")
            try:
                WebDriverWait(driver, 15).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "css-1djbb1k"))
                )
                jobs = find_elements_safely(driver, By.CLASS_NAME, "css-1djbb1k")
                if not jobs:
                    print(f"No jobs found for {role} on page {st.session_state.current_page}.")
                    break
                print(f"Found {len(jobs)} job listings for {role} on page {st.session_state.current_page}.")
                job_urls = [job.get_attribute("href") for job in jobs if job.get_attribute("href")]
                for apply_link in job_urls:
                    print(f"Scraping job details from: {apply_link}")
                    job_details = get_job_details(driver, apply_link)
                    job_details["Apply Link"] = apply_link
                    job_description = extract_job_description(driver)
                    job_details["Job Description"] = job_description.get("text", "Not available")
                    job_details["Job Description HTML"] = job_description.get("html", "")
                    structured_content = extract_structured_job_content(driver)
                    job_details["Structured Content"] = structured_content
                    results.append(job_details)
                    time.sleep(random.uniform(2, 4))
                next_buttons = find_elements_safely(driver, By.CLASS_NAME, "css-1u9zohc")
                if next_buttons and job_count < max_jobs and current_page < max_pages:
                    driver.execute_script("arguments[0].click();", next_buttons[0])
                    time.sleep(random.uniform(5, 8))
                    st.session_state.current_page += 1
                else:
                    break
            except Exception as e:
                print(f"Error processing page for {role}: {e}")
                break
    except Exception as e:
        print(f"Error in scrape_single_job: {e}")
    return results

#Function to scrape jobs and save all data into a single JSON file
def scrape_jobs_and_save(skills, driver):
    """
    Scrape jobs for the given skills and save the results into a single JSON file.
    """
    all_results = {}
    try:
        for skill in skills:
            st.write(f"🔍 Scraping jobs for skill: **{skill}**")
            print(f"Scraping jobs for skill: {skill}")
            
            # Scrape job listings for the skill
            skill_results = scrape_single_job(skill, driver)
            all_results[skill] = skill_results

        # Save all results into a single JSON file in the local directory
        file_name = "/home/muhammad-daniyal/Documents/POC/scraped_jobs.json"
        with open(file_name, "w", encoding="utf-8") as json_file:
            json.dump(all_results, json_file, indent=4)
        st.success(f"Job scraping completed. Results saved to {file_name}.")
    except Exception as e:
        st.error(f"An error occurred during the scraping process: {e}")
    finally:
        if 'driver' in locals() and driver:
            driver.quit()
            print("WebDriver closed successfully.")



# Initialize session state variable for profile and criteria
if "profile_and_criteria" not in st.session_state:
    st.session_state.profile_and_criteria = None

# ...existing code...

if options == "Motivation Matrix":
    st.title("Motivation Matrix")
    
    # Step 1: Purpose & Impact
    st.subheader("Purpose & Impact")
    purpose = st.text_area("Enter your purpose statement:", "")
    st.slider("Weight: Purpose & Impact", min_value=0, max_value=50, value=40, disabled=True)
    
    # Step 2: Strengths
    st.subheader("Strengths")
    strengths = st.text_area("Enter your top strengths:", "")
    st.slider("Weight: Strengths", min_value=0, max_value=50, value=25, disabled=True)
    
    # Step 3: Passions
    st.subheader("Passions")
    passions = st.text_area("Enter your passions:", "")
    st.slider("Weight: Passions", min_value=0, max_value=50, value=15, disabled=True)
    
    # Step 4: Motivations
    st.subheader("Motivations")
    motivations = st.text_area("Enter your motivations:", "")
    st.slider("Weight: Motivations", min_value=0, max_value=50, value=10, disabled=True)
    
    # Step 5: Priorities (Multi-answer section)
    st.subheader("Priorities")
    priorities = []
    for i in range(1, 6):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            priority = st.text_input(f"Priority {i}", key=f"priority_{i}")
        with col2:
            want = st.text_input(f"I want... {i}", key=f"want_{i}")
        with col3:
            nice_to_have = st.text_input(f"Would be nice to have... {i}", key=f"nice_to_have_{i}")
        with col4:
            dont_want = st.text_input(f"Don't Want {i}", key=f"dont_want_{i}")
        priorities.append({
            "Priority": priority,
            "I want": want,
            "Would be nice to have": nice_to_have,
            "Don't Want": dont_want
        })
    
    st.slider("Weight: Priorities", min_value=0, max_value=50, value=10, disabled=True)
    
    # Save Profile & Criteria button with unique key
    if st.button("Save Profile & Criteria", key="save_profile_criteria"):
        st.session_state.profile_and_criteria = {
            "Purpose": purpose,
            "Strengths": strengths,
            "Passions": passions,
            "Motivations": motivations,
            "Priorities": priorities
        }
        st.success("Profile and criteria saved temporarily in memory.")

    # Display saved profile and criteria if available
    if st.session_state.profile_and_criteria:
        st.markdown("### Saved Profile & Criteria")
        st.json(st.session_state.profile_and_criteria)

# ...existing code...

def scrape_jobs_by_level_and_location(skills, driver, job_level, job_location):
    """
    Scrape jobs based on skills, job level, and location.
    Respects the user-selected job and page limits.
    """
    results = {}
    max_jobs = st.session_state.get("num_jobs_per_skill", 10)  # Get user-selected job limit
    max_pages = st.session_state.get("num_pages_to_scrape", 5)  # Get user-selected page limit
    
    for skill in skills:
        st.write(f"🔍 Scraping jobs for skill: **{skill}**, Level: **{job_level}**, Location: **{job_location}**")
        
        # Initialize counters
        job_count = 0
        current_page = 1
        skill_results = []
        
        while job_count < max_jobs and current_page <= max_pages:
            # Build search URL with skill, level, and location
            level_term = "" if job_level == "Any Level" else f"+{job_level.replace(' ', '+')}"
            search_url = f"https://www.simplyhired.com/search?q={skill.replace(' ', '+')}{level_term}&l={job_location.replace(' ', '+')}&page={current_page}"
            
            try:
                # Navigate to the search page
                driver.get(search_url)
                time.sleep(random.uniform(5, 8))  # Prevent bot detection
                
                # Wait for job listings to load
                try:
                    WebDriverWait(driver, 15).until(
                        EC.presence_of_element_located((By.CLASS_NAME, "css-1djbb1k"))
                    )
                except:
                    st.warning(f"No jobs found for skill: {skill}, Level: {job_level}, Location: {job_location} on page {current_page}.")
                    break
                
                # Find all job listings
                jobs = find_elements_safely(driver, By.CLASS_NAME, "css-1djbb1k")
                if not jobs:
                    st.warning(f"No jobs found for skill: {skill}, Level: {job_level}, Location: {job_location} on page {current_page}.")
                    break
                
                # Extract job URLs
                job_urls = [job.get_attribute("href") for job in jobs if job.get_attribute("href")]
                
                # Process each job until we reach the maximum
                for apply_link in job_urls:
                    # Check if we've reached the job limit
                    if job_count >= max_jobs:
                        break
                        
                    # Status update
                    st.write(f"⏳ Processing job {job_count + 1}/{max_jobs} for skill: {skill}")
                    
                    # Extract job details
                    job_details = get_job_details(driver, apply_link)
                    job_details["Apply Link"] = apply_link
                    job_details["Skill"] = skill
                    job_details["Job Level"] = job_level
                    job_details["Job Location"] = job_location
                    
                    # Add to results and increment counter
                    skill_results.append(job_details)
                    job_count += 1
                    
                    # Progress update
                    st.write(f"✅ Completed {job_count}/{max_jobs} jobs for skill: {skill}")
                    
                    time.sleep(random.uniform(2, 4))
                    
                    # Check if we've reached the job limit
                    if job_count >= max_jobs:
                        break
                
                # Go to next page if needed
                if job_count < max_jobs and current_page < max_pages:
                    next_buttons = find_elements_safely(driver, By.CLASS_NAME, "css-1u9zohc")
                    if next_buttons:
                        driver.execute_script("arguments[0].click();", next_buttons[0])
                        time.sleep(random.uniform(5, 8))
                        current_page += 1
                    else:
                        st.write(f"No more pages available for skill: {skill}")
                        break
                else:
                    break
                    
            except Exception as e:
                st.error(f"Error processing page {current_page} for skill: {skill}. Error: {e}")
                break
        
        # Store results for this skill
        results[skill] = skill_results
        
        # Summary for this skill
        st.write(f"📊 Found {len(skill_results)} jobs for skill: **{skill}**")
    
    return results


if options == "View Scraped Jobs":
    st.markdown('<div class="main-title">View Scraped Job Results</div>', unsafe_allow_html=True)
    num_jobs_per_skill = st.sidebar.slider("Select number of jobs to scrape per skill", min_value=1, max_value=50, value=10)
    num_pages_to_scrape = st.sidebar.slider("Select number of pages to scrape", min_value=1, max_value=10, value=5)
    st.session_state.num_jobs_per_skill = num_jobs_per_skill
    st.session_state.num_pages_to_scrape = num_pages_to_scrape
    
    if st.session_state.cv_data is None:
        st.warning("No CV/Resume uploaded. Please upload a CV/Resume first.")
    else:
        # Step 1: Job Level Dropdown
        st.subheader("Select Job Level")
        job_level = st.selectbox(
            "Choose the job level you are looking for:",
            ["Any Level", "Entry Level", "Junior Level", "Senior Level", "Expert Level"]
        )
        
        # Step 2: Job Location Input
        st.subheader("Enter Job Location")
        job_location = st.text_input("Enter the location where you want the job (e.g., New York, Remote):", "")
        
        # Step 3: Scraping Button
        if st.button("Start Scraping Jobs"):
            if not job_location.strip():
                st.warning("Please enter a job location.")
            else:
                st.session_state.scraping_in_progress = True
                try:
                    driver = init_driver()
                    with st.spinner(f"Scraping jobs for level: {job_level}, location: {job_location}..."):
                        skills = st.session_state.cv_data.get("Skills", [])
                        if skills and skills != ["Skills not found"]:
                            job_results = scrape_jobs_by_level_and_location(skills, driver, job_level, job_location)
                            st.session_state.job_results = job_results
                            # save_to_json(job_results)
                            st.success("Job scraping completed. Results saved.")
                        else:
                            st.warning("No skills found in the CV/Resume. Job scraping skipped.")
                except Exception as e:
                    st.error(f"An error occurred during job scraping: {e}")
                finally:
                    if 'driver' in locals() and driver:
                        driver.quit()
                    st.session_state.scraping_in_progress = False
        
        # # Display job results if available
        # if st.session_state.job_results:
        #     if "display_job_results" in globals():
        #         display_job_results(st.session_state.job_results)
        #     else:
        #         st.error("The function 'display_job_results' is not defined.")
        # else:
        #     st.warning("No job results to display. Please start the scraping process.")


def display_job_results(job_results):
    """
    Display job scraping results in Streamlit, including salary, job type, job location, job posted, and structured data.
    """
    st.markdown("### Scraped Job Results")

    if not job_results:
        st.warning("No job results to display.")
        return

    # Count total jobs
    total_jobs = sum(len(jobs) for jobs in job_results.values())
    st.write(f"### 🔍 Found {total_jobs} jobs across {len(job_results)} skills")
    
    # Global slider for maximum jobs to display per skill
    max_jobs_per_skill = st.slider("Maximum jobs to display per skill", 
                                   min_value=5, 
                                   max_value=50, 
                                   value=20, 
                                   step=5)

    # Flatten the job results for display
    flattened_results = []
    for skill, jobs in job_results.items():
        jobs_to_display = jobs[:max_jobs_per_skill]  # Apply the slider limit
        
        for job in jobs_to_display:
            if isinstance(job, dict):  # Ensure job is a dictionary
                # Verify the job title or description contains the exact skill keyword
                job_text = (job.get("Job Title", "").lower() + " " + job.get("Job Description", "").lower())
                skill_pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if not re.search(skill_pattern, job_text):
                    print(f"Skipping job '{job.get('Job Title', 'Unknown')}' - doesn't match skill '{skill}'")
                    continue

                flattened_results.append(job)

    # Create tabs for different skills
    tabs = st.tabs(list(job_results.keys()))
    
    # Fill each tab with job data
    for i, (skill, jobs) in enumerate(job_results.items()):
        with tabs[i]:
            st.write(f"### {skill}: {len(jobs)} jobs found")
            
            if not jobs:
                st.warning(f"No jobs found for {skill}.")
                continue
            
            # Apply the jobs limit from the slider
            jobs_to_display = jobs[:max_jobs_per_skill]
            
            if len(jobs) > max_jobs_per_skill:
                st.info(f"Showing {max_jobs_per_skill} of {len(jobs)} jobs. Adjust the slider above to see more.")
            
            # Create an expander for each job
            for j, job in enumerate(jobs_to_display):
                if not isinstance(job, dict):
                    continue
                    
                with st.expander(f"{j+1}. {job.get('Job Title', 'Unknown Title')} - {job.get('Company Name', 'Unknown Company')}"):
                    # Job header
                    st.markdown(f"**🏢 Company:** {job.get('Company Name', 'Not available')}")
                    st.markdown(f"**📍 Location:** {job.get('Job Location', 'Not available')}")
                    st.markdown(f"**💰 Salary:** {job.get('Salary', 'Not disclosed')}")
                    st.markdown(f"**🕒 Posted:** {job.get('Posted', 'Not available')}")
                    st.markdown(f"**👔 Job Type:** {job.get('Job Type', 'Not specified')}")
                    
                    # Job description
                    st.markdown("### Job Description")
                    st.markdown(job.get('Job Description', 'Not available'))
                    
                    # Display HTML version of job description if available
                    if job.get('Job Description HTML'):
                        st.components.v1.html(job.get('Job Description HTML', ''), height=500)
                    
                    # Benefits and Qualifications
                    if job.get('Benefits') and job.get('Benefits') != 'Not specified':
                        st.markdown("### Benefits")
                        st.markdown(job.get('Benefits'))
                        
                    if job.get('Qualifications') and job.get('Qualifications') != 'Not specified':
                        st.markdown("### Qualifications")
                        st.markdown(job.get('Qualifications'))
                    
                    # Apply link
                    st.markdown("### Apply")
                    if job.get('Apply Link'):
                        st.markdown(f"[Apply Now]({job.get('Apply Link')})")
                    else:
                        st.warning("No application link available.")

def main():
    # Initialize session state variables if they don't exist
    if 'uploaded_file' not in st.session_state:
        st.session_state.uploaded_file = None
    if 'uploaded_file_content' not in st.session_state:
        st.session_state.uploaded_file_content = None
    if 'cv_data' not in st.session_state:
        st.session_state.cv_data = None
    if 'job_results' not in st.session_state:
        st.session_state.job_results = None
    if 'scraping_in_progress' not in st.session_state:
        st.session_state.scraping_in_progress = False
    if 'process_completed' not in st.session_state:
        st.session_state.process_completed = False

    # Set up the main page layout
    st.title("🎯 DREAM JOB CATCHER")
    st.markdown("""
        <style>
            .main-title {
                font-size: 24px;
                font-weight: bold;
                color: #34495e;
                margin-top: 20px;
            }
            .section-title {
                font-size: 20px;
                font-weight: bold;
                color: #2980b9;
                margin-top: 15px;
            }
            .info-box {
                background-color: black;
                padding: 10px;
                border-radius: 5px;
                margin-bottom: 10px;
            }
            .sub-title {
                font-size: 18px;
                font-weight: bold;
                color: #3498db;
                margin-top: 10px;
            }
        </style>
    """, unsafe_allow_html=True)

    # Sidebar for navigation
    st.sidebar.title("Navigation")
    st.sidebar.markdown("Use the options below to navigate:")
    options = st.sidebar.radio("Choose an option:", ["Upload CV/Resume", "View Scraped Jobs", "Motivation Matrix", "Job Ranking & Score", "About"])

    if options == "Upload CV/Resume":
        st.markdown('<div class="main-title">Upload Your CV/Resume</div>', unsafe_allow_html=True)
        
        # Check if a file is already uploaded in session state
        if st.session_state.uploaded_file is None:
            uploaded_file = st.file_uploader(
                "Upload a CV/Resume", 
                type=["pdf", "docx", "txt"]
            )
            if uploaded_file is not None:
                # Save the uploaded file in session state
                st.session_state.uploaded_file = uploaded_file
                
                # Create a tempfile to store the uploaded file
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    temp_file_path = tmp_file.name
                
                try:
                    # Extract text based on file type
                    if uploaded_file.name.endswith('.pdf'):
                        text = extract_text_from_pdf(temp_file_path)
                        file_type = "PDF"
                    elif uploaded_file.name.endswith('.docx'):
                        text = extract_text_from_docx(temp_file_path)
                        file_type = "DOCX"
                    else:  # txt
                        text = uploaded_file.getvalue().decode("utf-8")
                        file_type = "TXT"
                    
                    # Save the extracted text in session state
                    st.session_state.uploaded_file_content = text
                    
                    # Delete temporary file
                    os.unlink(temp_file_path)
                    
                    # Display success message
                    st.success(f"Successfully processed {file_type} file")
                    
                except Exception as e:
                    st.error(f"An error occurred: {e}")
        else:
            # Display the uploaded file name
            st.info(f"Uploaded file: {st.session_state.uploaded_file.name}")
        
        # Perform CV analysis if content is available
        if st.session_state.uploaded_file_content is not None:
            st.markdown('<div class="sub-title">Extracted CV Data</div>', unsafe_allow_html=True)
            st.session_state.cv_data = perform_cv_analysis(st.session_state.uploaded_file_content)
        
        # Forte Document Extractor integration
        st.markdown('<div class="sub-title"></div>', unsafe_allow_html=True)
        if st.session_state.uploaded_file is not None:
            # Save the uploaded file to a temporary directory
            temp_dir = "temp"
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, st.session_state.uploaded_file.name)
            
            with open(temp_file_path, "wb") as f:
                f.write(st.session_state.uploaded_file.getbuffer())
            
            # Initialize the extractor
            extractor = ForteDocumentExtractor()  # Ensure this class is defined or imported
            
           
        
        # Add a separate button for uploading Forte Document
        st.markdown('<div class="sub-title">Upload Forte Document</div>', unsafe_allow_html=True)
        forte_uploaded_file = st.file_uploader(
            "Upload Forte Document (PDF, DOCX, DOC, or TXT)", 
            type=["pdf", "docx", "doc", "txt"], 
            key="forte_file_uploader"
        )
        if forte_uploaded_file is not None:
            # Ensure temp_dir is initialized
            temp_dir = "temp"
            os.makedirs(temp_dir, exist_ok=True)
            
            # Save the uploaded Forte file to a temporary directory
            forte_temp_file_path = os.path.join(temp_dir, forte_uploaded_file.name)
            with open(forte_temp_file_path, "wb") as f:
                f.write(forte_uploaded_file.getbuffer())
            
            st.success(f"Forte Document uploaded: {forte_uploaded_file.name}")
            
            # Initialize the extractor
            extractor = ForteDocumentExtractor()  # Ensure this class is defined or imported
            
            # Process the Forte document
            if st.button("Extract Forte Document Data"):
                with st.spinner("Extracting data from the Forte document..."):
                    forte_result = extractor.extract_document(forte_temp_file_path)
                
                if "error" in forte_result:
                    st.error(f"Error extracting Forte document data: {forte_result['error']}")
                else:
                    st.success("Forte Document data extracted successfully!")
                    st.json(forte_result)

    elif options == "View Scraped Jobs":
        st.markdown('<div class="main-title">View Scraped Job Results</div>', unsafe_allow_html=True)
        num_jobs_per_skill = st.slider("Select number of jobs to scrape per skill", min_value=1, max_value=50, value=10)
        num_pages_to_scrape = st.slider("Select number of pages to scrape", min_value=1, max_value=10, value=5)
        st.session_state.num_jobs_per_skill = num_jobs_per_skill
        st.session_state.num_pages_to_scrape = num_pages_to_scrape
        
        # Ensure CV data is available before starting scraping
        if st.session_state.cv_data is None:
            st.warning("No CV/Resume uploaded. Please upload a CV/Resume first.")
            return
        
        # Step 1: Job Level Dropdown
        st.subheader("Select Job Level")
        job_level = st.selectbox(
            "Choose the job level you are looking for:",
            ["Any Level", "Entry Level", "Junior Level", "Senior Level", "Expert Level"]    
        )
        
        # Step 2: Job Location Input
        st.subheader("Enter Job Location")
        job_location = st.text_input("Enter the location where you want the job (e.g., New York, Remote):", "")
        
        # Step 3: Scraping Button
        if st.button("Start Scraping Jobs"):
            if not job_location.strip():
                st.warning("Please enter a job location.")
            else:
                st.session_state.scraping_in_progress = True
                try:
                    driver = init_driver()
                    with st.spinner(f"Scraping jobs for level: {job_level}, location: {job_location}..."):
                        skills = st.session_state.cv_data.get("Skills", [])
                        if skills and skills != ["Skills not found"]:
                            job_results = scrape_jobs_by_level_and_location(skills, driver, job_level, job_location)
                            st.session_state.job_results = job_results
                            st.success("Job scraping completed. Results saved.")
                        else:
                            st.warning("No skills found in the CV/Resume. Job scraping skipped.")
                except Exception as e:
                    st.error(f"An error occurred during job scraping: {e}")
                finally:
                    if 'driver' in locals() and driver:
                        driver.quit()
                    st.session_state.scraping_in_progress = False
        
        # Display job results if available
        if st.session_state.job_results:
            display_job_results(st.session_state.job_results)
        else:
            st.warning("No job results to display. Please start the scraping process.")

    elif options == "Motivation Matrix":
        st.title("Motivation Matrix")
        
        # Step 1: Purpose & Impact
        st.subheader("Purpose & Impact")
        purpose = st.text_area("Enter your purpose statement:", "")
        st.slider("Weight: Purpose & Impact", min_value=0, max_value=50, value=40, disabled=True)
        
        # Step 2: Strengths
        st.subheader("Strengths")
        strengths = st.text_area("Enter your top strengths:", "")
        st.slider("Weight: Strengths", min_value=0, max_value=50, value=25, disabled=True)
        
        # Step 3: Passions
        st.subheader("Passions")
        passions = st.text_area("Enter your passions:", "")
        st.slider("Weight: Passions", min_value=0, max_value=50, value=15, disabled=True)
        
        # Step 4: Motivations
        st.subheader("Motivations")
        motivations = st.text_area("Enter your motivations:", "")
        st.slider("Weight: Motivations", min_value=0, max_value=50, value=10, disabled=True)
        
        # Step 5: Priorities (Multi-answer section)
        st.subheader("Priorities")
        priorities = []
        for i in range(1, 6):
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                priority = st.text_input(f"Priority {i}", key=f"priority_{i}")
            with col2:
                want = st.text_input(f"I want... {i}", key=f"want_{i}")
            with col3:
                nice_to_have = st.text_input(f"Would be nice to have... {i}", key=f"nice_to_have_{i}")
            with col4:
                dont_want = st.text_input(f"Don't Want {i}", key=f"dont_want_{i}")
            priorities.append({
                "Priority": priority,
                "I want": want,
                "Would be nice to have": nice_to_have,
                "Don't Want": dont_want
            })
    
        st.slider("Weight: Priorities", min_value=0, max_value=50, value=10, disabled=True)
        
        # Save Profile & Criteria button with unique key
        if st.button("Save Profile & Criteria", key="save_profile_criteria"):
            st.session_state.profile_and_criteria = {
                "Purpose": purpose,
                "Strengths": strengths,
                "Passions": passions,
                "Motivations": motivations,
                "Priorities": priorities
            }
            st.success("Profile and criteria saved temporarily in memory.")

        # Display saved profile and criteria if available
        if st.session_state.profile_and_criteria:
            st.markdown("### Saved Profile & Criteria")
            st.json(st.session_state.profile_and_criteria)

    elif options == "About":
        st.markdown('<div class="main-title">About This Application</div>', unsafe_allow_html=True)
        st.markdown("""
            This application allows users to:
            - Upload their CV/Resume for analysis.
            - Extract key details such as skills, experience, and qualifications.
            - Scrape job listings based on the extracted skills.
            - View and download the scraped job results.
            
            **Developed by:** Muhammad Daniyal  
            **Technologies Used:** Python, Streamlit, Selenium, Plotly, Pandas
        """)

    elif options == "Job Ranking & Score":
        st.markdown('<div class="main-title">Job Ranking & Score</div>', unsafe_allow_html=True)
        
        # Create a clean layout with columns
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("""
            <div style="background-color:#f8f9fa; padding:15px; border-radius:10px; margin-bottom:20px">
                <h3 style="color:#1f77b4; margin-top:0">📊 Job Ranking Analysis</h3>
                <p>This tool analyzes job listings against your CV and motivation criteria to determine the best matches for your career goals.</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.image("https://img.icons8.com/fluency/96/000000/job-seeker.png", width=100)
        
        # Analysis button with enhanced styling
        analyze_button = st.button("Analyze Job Rankings", use_container_width=True)
        
        if analyze_button:
            with st.spinner("📊 Analyzing and ranking jobs based on your profile..."):
                try:
                    # Initialize OpenAI conversation
                    conversation = initialize_openai_chat_completion()
                    
                    # Construct the prompt
                    prompt = construct_job_ranking_prompt(
                        job_results=st.session_state.job_results,
                        cv_data=st.session_state.cv_data,
                        motivation_matrix=st.session_state.profile_and_criteria
                    )
                    
                    # Add a toggle to show/hide the prompt
                    with st.expander("📜 Show Prompt Details", expanded=False):
                        st.text_area("Final Prompt Sent to OpenAI:", prompt, height=250)
                    
                    # Call OpenAI API
                    ranking_response = get_openai_job_ranking(conversation, prompt)
                    
                    # Add a toggle to show/hide the raw response
                    with st.expander("🔍 Show Raw API Response", expanded=False):
                        st.text_area("OpenAI Response:", ranking_response, height=200)
                    
                    # # Handle and parse OpenAI response
                    # if not ranking_response:
                    #     st.error("⚠️ OpenAI did not return a valid response.")
                    #     st.stop()
                    
                    # try:
                    #     ranking_data = json.loads(ranking_response)
                    # except json.JSONDecodeError:
                    #     st.error("⚠️ OpenAI returned an invalid JSON response.")
                    #     st.stop()

                    
                    
                    # # Display ranked job results with visualization
                    # st.markdown("<h2 style='color:#1f77b4; margin-top:30px'>📊 Job Match Results</h2>", unsafe_allow_html=True)
                    
                    # # Extract rankings for visualization
                    # job_titles = []
                    # scores = []
                    # rankings = []
                    
                    # # Process the ranking data
                    # for job in ranking_data["rankings"]:
                    #     job_titles.append(job.get("job_title", "Unknown Job"))
                    #     scores.append(job.get("score", 0))
                    #     rankings.append(job.get("ranking", 0))
                    
                    # # Create tabs for different visualizations
                    # tab1, tab2, tab3 = st.tabs(["Score Comparison", "Ranking Chart", "Detailed Table"])
                    
                    # with tab1:
                    #     # Create a horizontal bar chart for scores
                    #     fig1 = go.Figure()
                    #     fig1.add_trace(go.Bar(
                    #         y=job_titles,
                    #         x=scores,
                    #         orientation='h',
                    #         marker=dict(
                    #             color=scores,
                    #             colorscale='Viridis',
                    #             colorbar=dict(title="Score %"),
                    #         ),
                    #         text=[f"{score}%" for score in scores],
                    #         textposition='auto',
                    #     ))
                        
                    #     fig1.update_layout(
                    #         title="Job Match Scores",
                    #         xaxis_title="Match Score (%)",
                    #         yaxis_title="Job Title",
                    #         height=400,
                    #         margin=dict(l=20, r=20, t=40, b=20),
                    #         xaxis=dict(range=[0, 100]),
                    #     )
                        
                    #     st.plotly_chart(fig1, use_container_width=True)
                    
                    # with tab2:
                    #     # Create a scatter plot for rankings
                    #     fig2 = go.Figure()
                    #     fig2.add_trace(go.Scatter(
                    #         x=rankings,
                    #         y=job_titles,
                    #         mode='markers',
                    #         marker=dict(
                    #             size=16,
                    #             color=rankings,
                    #             colorscale='Viridis',
                    #             showscale=True,
                    #             colorbar=dict(title="Ranking"),
                    #             reversescale=True,
                    #         ),
                    #         text=[f"Rank: {rank}/10" for rank in rankings],
                    #         hoverinfo="text+y"
                    #     ))
                        
                    #     fig2.update_layout(
                    #         title="Job Rankings (1-10 scale)",
                    #         xaxis_title="Ranking Score",
                    #         yaxis_title="Job Title",
                    #         height=400,
                    #         margin=dict(l=20, r=20, t=40, b=20),
                    #         xaxis=dict(range=[0, 11]),
                    #     )
                        
                    #     st.plotly_chart(fig2, use_container_width=True)
                    
                    # with tab3:
                    #     # Create a detailed table with all information
                    #     st.markdown("### Detailed Job Rankings")
                        
                    #     # Create a styled DataFrame for better visualization
                    #     df = pd.DataFrame(ranking_data["rankings"])
                        
                    #     # Add conditional formatting based on score
                    #     def color_score(val):
                    #         if isinstance(val, (int, float)):
                    #             if val >= 80:
                    #                 return f'background-color: #c6efce; color: #006100'
                    #             elif val >= 60:
                    #                 return f'background-color: #ffeb9c; color: #9c5700'
                    #             else:
                    #                 return f'background-color: #ffc7ce; color: #9c0006'
                    #         return ''
                        
                    #     # Style the DataFrame
                    #     styled_df = df.style.applymap(color_score, subset=['score'])
                        
                    #     # Display the table
                    #     st.dataframe(styled_df, use_container_width=True)
                    
                    # # Add explanation section
                    # st.markdown("""
                    # <div style="background-color:#f8f9fa; padding:15px; border-radius:10px; margin-top:30px">
                    #     <h3 style="color:#1f77b4; margin-top:0">📌 Understanding Your Results</h3>
                    #     <p><strong>Score:</strong> Percentage match between your profile and the job (0-100%)</p>
                    #     <p><strong>Ranking:</strong> Overall job suitability rating on a scale of 1-10</p>
                    #     <p>Jobs with higher scores and rankings are better aligned with your skills, experience, and motivation criteria.</p>
                    # </div>
                    # """, unsafe_allow_html=True)
                    
                    # # Export options
                    # st.markdown("<h3 style='color:#1f77b4; margin-top:30px'>🔄 Export Options</h3>", unsafe_allow_html=True)
                    
                    # col1, col2 = st.columns(2)
                    # with col1:
                    #     if st.button("Export as CSV", use_container_width=True):
                    #         csv = df.to_csv(index=False)
                    #         st.download_button(
                    #             label="Download CSV",
                    #             data=csv,
                    #             file_name="job_rankings.csv",
                    #             mime="text/csv",
                    #             use_container_width=True
                    #         )
                    
                    # with col2:
                    #     if st.button("Export as JSON", use_container_width=True):
                    #         json_str = df.to_json(orient="records")
                    #         st.download_button(
                    #             label="Download JSON",
                    #             data=json_str,
                    #             file_name="job_rankings.json",
                    #             mime="application/json",
                    #             use_container_width=True
                    #         )
                
                except Exception as e:
                    st.error(f"An error occurred during job ranking: {e}")
                    logger.exception("Exception during job ranking:")

def initialize_openai_chat_completion():
    """
    Initialize the conversation for ChatCompletion using a detailed system prompt.
    """
    system_message = {
        "role": "system",
        "content": (
            "You are an expert job ranking assistant. You will be provided with a candidate's CV data, "
            "their Motivation Matrix inputs, and detailed scraped job information. Your task is to evaluate "
            "each job and rank it on a scale of 1 to 10, then calculate a job score percentage (0-100) based on "
            "how well the job aligns with the candidate's requirements. The candidate's preferences have the following weights: "
            "Purpose & Impact: 40, Strengths: 25, Passions: 15, Motivations: 10, Priorities: 10. "
            "Please provide your output in valid JSON format, listing each job with its ranking and score."
        )
    }
    conversation = [system_message]
    logger.info("ChatCompletion conversation initialized.")
    return conversation

def construct_job_ranking_prompt(job_results, cv_data, motivation_matrix):
    """
    Construct a detailed prompt string combining scraped job details, CV data, and the Motivation Matrix.
    """
    prompt_lines = []
    
    # Section 1: Candidate CV Data
    prompt_lines.append("=== Candidate CV Data ===")
    prompt_lines.append(json.dumps(cv_data, indent=2))
    
    # Section 2: Candidate Motivation Matrix
    prompt_lines.append("\n=== Candidate Motivation Matrix ===")
    prompt_lines.append(json.dumps(motivation_matrix, indent=2))
    
    # Section 3: Detailed Scraped Job Listings
    prompt_lines.append("\n=== Scraped Job Details ===")
    for skill, jobs in job_results.items():
        prompt_lines.append(f"\n--- Skill: {skill} ---")
        for idx, job in enumerate(jobs, start=1):
            job_text = (
                f"Job #{idx}:\n"
                f"  Job Title       : {job.get('Job Title', 'N/A')}\n"
                f"  Company Name    : {job.get('Company Name', 'N/A')}\n"
                f"  Job Location    : {job.get('Job Location', 'N/A')}\n"
                f"  Salary          : {job.get('Salary', 'N/A')}\n"
                f"  Job Type        : {job.get('Job Type', 'N/A')}\n"
                f"  Job Rating      : {job.get('Job Rating', 'N/A')}\n"
                f"  Posted          : {job.get('Posted', 'N/A')}\n"
                f"  Benefits        : {job.get('Benefits', 'N/A')}\n"
                f"  Qualifications  : {job.get('Qualifications', 'N/A')}\n"
                f"  Job Description : {job.get('Job Description', 'N/A')}\n"
                f"  Apply Link      : {job.get('Apply Link', 'N/A')}\n"
                f"  Structured Data : {job.get('Structured Content', 'N/A')}\n"
            )
            prompt_lines.append(job_text)
    
    # Final instruction
    prompt_lines.append(
        "\nBased on the above data, evaluate and rank each job on a scale from 1 to 10 and compute a job score percentage (0-100) "
        "that indicates how closely each job aligns with the candidate's profile and motivation. Provide your response in valid JSON format."
    )
    prompt = "\n".join(prompt_lines)
    logger.info("Constructed detailed job ranking prompt.")
    return prompt



# Retrieve the API key from environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Function to handle OpenAI API errors gracefully
def get_openai_job_ranking(conversation, prompt):
    """
    Sends a chat completion request to OpenAI's API for job ranking.
    Ensures the API key is valid and handles errors properly.
    """
    if not OPENAI_API_KEY:
        raise ValueError("Missing OpenAI API key. Please set OPENAI_API_KEY as an environment variable.")

    # Append the user prompt to the conversation history
    conversation.append({"role": "user", "content": prompt})

    try:
        # OpenAI API call using the latest correct format
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o",  # Ensure this model is accessible in your OpenAI account
            messages=conversation,
            temperature=0.7,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            n=1,
            response_format={"type": "json_object"}  # Ensure the response is structured as JSON
        )

        # Extract response content
        if response.choices:
            return response.choices[0].message.content  # Extracts the assistant's response
        else:
            return None  # No response received

    except openai.AuthenticationError:
        print("⚠️ OpenAI API Authentication Failed: Check your API key.")
        return None
    except openai.RateLimitError:
        print("⚠️ Rate Limit Exceeded: Too many requests sent to OpenAI API.")
        return None
    except openai.OpenAIError as e:
        print(f"⚠️ OpenAI API Error: {e}")
        return None
    except Exception as e:
        print(f"⚠️ Unexpected Error: {e}")
        return None
    

def display_job_results(job_results):
    """
    Display job scraping results in a professional Streamlit interface with advanced filtering,
    visualization, and organized presentation.
    """


    # Professional header with styling
    st.markdown("""
    # <div style='text-align: center; background-color: #f0f2f6; padding: 20px; border-radius: 10px;'>
        <h1 style='color: #1E3A8A;'>Job Search Dashboard</h1>
        <p style='font-size: 20px;'>Find your next career opportunity</p>
    </div>
    """, unsafe_allow_html=True)

    if not job_results:
        st.warning("No job results to display.")
        return

    # Count total jobs
    total_jobs = sum(len(jobs) for jobs in job_results.values())
    
    # Dashboard metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Jobs Found", total_jobs)
    with col2:
        st.metric("Skills Searched", len(job_results))
    with col3:
        st.metric("Last Updated", datetime.now().strftime("%Y-%m-%d %H:%M"))

    # Sidebar for filters
    st.sidebar.markdown("## Job Filters")
    max_jobs_per_skill = st.sidebar.slider(
        "Maximum jobs per skill", 
        min_value=5, 
        max_value=50, 
        value=20, 
        step=5
    )
    
    # Create a unified job list for filtering
    all_jobs = []
    for skill, jobs in job_results.items():
        for job in jobs:
            if isinstance(job, dict):
                job['Skill'] = skill
                all_jobs.append(job)
    
    # Extract unique job types and locations for filtering
    if all_jobs:
        job_types = ["All"] + list(set(job.get('Job Type', 'Not specified') for job in all_jobs if isinstance(job.get('Job Type'), str)))
        job_locations = ["All"] + list(set(job.get('Job Location', 'Not specified') for job in all_jobs if isinstance(job.get('Job Location'), str)))
        
        # Add filters to sidebar
        job_type_filter = st.sidebar.selectbox("Job Type", job_types)
        job_location_filter = st.sidebar.selectbox("Location", job_locations)
        
        # Salary range filter
        st.sidebar.markdown("### Salary Filter")
        show_salary_only = st.sidebar.checkbox("Show only jobs with salary info")
        
        # Apply filters
        filtered_jobs = all_jobs
        if job_type_filter != "All":
            filtered_jobs = [job for job in filtered_jobs if job.get('Job Type') == job_type_filter]
        if job_location_filter != "All":
            filtered_jobs = [job for job in filtered_jobs if job.get('Job Location') == job_location_filter]
        if show_salary_only:
            filtered_jobs = [job for job in filtered_jobs if job.get('Salary') and job.get('Salary') != 'Not disclosed']
        
        # Rebuild job_results with filtered data
        filtered_results = {}
        for job in filtered_jobs:
            skill = job['Skill']
            if skill not in filtered_results:
                filtered_results[skill] = []
            filtered_results[skill].append(job)
    else:
        filtered_results = job_results

    # Data visualization
    st.markdown("## Job Market Overview")
    viz_col1, viz_col2 = st.columns(2)
    
    if all_jobs:
        # Job distribution by skill
        skill_counts = {}
        for skill, jobs in filtered_results.items():
            skill_counts[skill] = len(jobs)
        
        with viz_col1:
            if skill_counts:
                fig = px.pie(
                    names=list(skill_counts.keys()),
                    values=list(skill_counts.values()),
                    title="Jobs by Skill",
                    color_discrete_sequence=px.colors.qualitative.Pastel
                )
                st.plotly_chart(fig, use_container_width=True)
        
        # Job types distribution
        job_type_counts = {}
        for job in all_jobs:
            job_type = job.get('Job Type', 'Not specified')
            if job_type not in job_type_counts:
                job_type_counts[job_type] = 0
            job_type_counts[job_type] += 1
        
        with viz_col2:
            if job_type_counts:
                fig = px.bar(
                    x=list(job_type_counts.keys()),
                    y=list(job_type_counts.values()),
                    title="Jobs by Type",
                    color=list(job_type_counts.keys()),
                    color_discrete_sequence=px.colors.qualitative.Pastel
                )
                st.plotly_chart(fig, use_container_width=True)

    # Main job listings
    st.markdown("## Job Listings")
    
    # Create tabs for different skills
    if filtered_results:
        tabs = st.tabs(list(filtered_results.keys()))
        
        for i, (skill, jobs) in enumerate(filtered_results.items()):
            with tabs[i]:
                st.write(f"### {skill}: {len(jobs)} jobs found")
                
                if not jobs:
                    st.warning(f"No jobs found for {skill} with the current filters.")
                    continue
                
                # Apply the jobs limit from the slider
                jobs_to_display = jobs[:max_jobs_per_skill]
                
                if len(jobs) > max_jobs_per_skill:
                    st.info(f"Showing {max_jobs_per_skill} of {len(jobs)} jobs. Adjust the slider to see more.")
                
                # Create job cards
                job_cards = []
                for j, job in enumerate(jobs_to_display):
                    if not isinstance(job, dict):
                        continue
                    
                    # Job card container
                    with st.container():
                        st.markdown(f"""
                        <div style='background-color: white; padding: 15px; border-radius: 10px; 
                                    margin-bottom: 15px; border: 1px solid #e0e0e0; box-shadow: 0 2px 5px rgba(0,0,0,0.1);'>
                            <h3 style='color: #1E3A8A;'>{job.get('Job Title', 'Unknown Title')}</h3>
                            <p style='font-size: 16px; color: #4B5563;'>{job.get('Company Name', 'Unknown Company')}</p>
                            <div style='display: flex; flex-wrap: wrap; gap: 10px; margin-top: 10px;'>
                                <span style='background-color: #E5E7EB; padding: 5px 10px; border-radius: 15px; font-size: 14px;'>
                                    📍 {job.get('Job Location', 'Not available')}
                                </span>
                                <span style='background-color: #E5E7EB; padding: 5px 10px; border-radius: 15px; font-size: 14px;'>
                                    💰 {job.get('Salary', 'Not disclosed')}
                                </span>
                                <span style='background-color: #E5E7EB; padding: 5px 10px; border-radius: 15px; font-size: 14px;'>
                                    👔 {job.get('Job Type', 'Not specified')}
                                </span>
                                <span style='background-color: #E5E7EB; padding: 5px 10px; border-radius: 15px; font-size: 14px;'>
                                    🕒 {job.get('Posted', 'Not available')}
                                </span>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Job details expander
                        with st.expander("View Details"):
                            tab1, tab2, tab3 = st.tabs(["Description", "Details", "Apply"])
                            
                            with tab1:
                                # Job description
                                st.markdown("### Job Description")
                                st.markdown(job.get('Job Description', 'Not available'))
                                
                                # Display HTML version of job description if available
                                if job.get('Job Description HTML'):
                                    st.markdown("### Full Description")
                                    st.components.v1.html(job.get('Job Description HTML', ''), height=500)
                            
                            with tab2:
                                # Benefits and Qualifications
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.markdown("### Benefits")
                                    if isinstance(job.get('Benefits'), list):
                                        for benefit in job.get('Benefits', ['Not specified']):
                                            st.markdown(f"- {benefit}")
                                    else:
                                        st.markdown(job.get('Benefits', 'Not specified'))
                                        
                                with col2:
                                    st.markdown("### Qualifications")
                                    if isinstance(job.get('Qualifications'), list):
                                        for qual in job.get('Qualifications', ['Not specified']):
                                            st.markdown(f"- {qual}")
                                    else:
                                        st.markdown(job.get('Qualifications', 'Not specified'))
                                
                                # Structured content
                                structured = job.get('Structured Content', {})
                                if structured and isinstance(structured, dict):
                                    st.markdown("### Details")
                                    
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        # Requirements
                                        if 'Requirements' in structured:
                                            st.markdown("#### Requirements")
                                            for req in structured['Requirements']:
                                                st.markdown(f"- {req}")
                                    
                                    with col2:
                                        # Responsibilities
                                        if 'Responsibilities' in structured:
                                            st.markdown("#### Responsibilities")
                                            for resp in structured['Responsibilities']:
                                                st.markdown(f"- {resp}")
                                    
                                    # Benefits
                                    if 'Benefits' in structured:
                                        st.markdown("#### Benefits")
                                        for benefit in structured['Benefits']:
                                            st.markdown(f"- {benefit}")
                                    
                                    # Any other structured content sections
                                    for key, values in structured.items():
                                        if key not in ['Requirements', 'Responsibilities', 'Benefits'] and isinstance(values, list):
                                            st.markdown(f"#### {key}")
                                            for item in values:
                                                st.markdown(f"- {item}")
                            
                            with tab3:
                                # Apply link
                                st.markdown("### Apply for this Position")
                                if job.get('Apply Link'):
                                    st.markdown(f"""
                                    <a href="{job.get('Apply Link')}" target="_blank">
                                        <button style="background-color: #1E3A8A; color: white; padding: 10px 20px; 
                                                border: none; border-radius: 5px; cursor: pointer; font-weight: bold;">
                                            Apply Now
                                        </button>
                                    </a>
                                    """, unsafe_allow_html=True)
                                else:
                                    st.warning("No application link available.")
    
    # Export options
    st.sidebar.markdown("## Export Options")
    if st.sidebar.button("Export to CSV"):
        if all_jobs:
            # Convert to DataFrame
            df = pd.DataFrame(all_jobs)
            # Create a download link
            csv = df.to_csv(index=False)
            st.sidebar.download_button(
                label="Download CSV",
                data=csv,
                file_name="job_search_results.csv",
                mime="text/csv"
            )
        else:
            st.sidebar.warning("No data to export")
    
    # Footer
    st.markdown("""
    <div style='text-align: center; margin-top: 30px; padding: 20px; border-top: 1px solid #e0e0e0;'>
        <p>Job Search Dashboard | Powered by Streamlit</p>
    </div>
    """, unsafe_allow_html=True)

# Example usage:
if __name__ == "__main__":
    # Sample data for testing
    sample_job_results = {
        "Python Developer": [
            {
                "Job Title": "Senior Python Developer",
                "Company Name": "Tech Solutions Inc.",
                "Job Location": "San Francisco, CA",
                "Salary": "$120,000 - $150,000",
                "Job Type": "Full-time",
                "Job Rating": "4.5/5",
                "Posted": "2 days ago",
                "Benefits": ["Health Insurance", "401k", "Remote Work"],
                "Qualifications": ["5+ years Python experience", "Django", "AWS"],
                "Job Description": "We are looking for an experienced Python developer to join our team...",
                "Apply Link": "https://example.com/apply",
                "Structured Content": {
                    "Requirements": ["Python expertise", "Web development skills", "Database knowledge"],
                    "Responsibilities": ["Develop new features", "Maintain existing code", "Collaborate with team"]
                }
            },
            # Add more sample jobs here
        ],
        "Data Scientist": [
            {
                "Job Title": "Senior Data Scientist",
                "Company Name": "Data Analytics Co.",
                "Job Location": "New York, NY",
                "Salary": "$130,000 - $160,000",
                "Job Type": "Full-time",
                "Job Rating": "4.2/5",
                "Posted": "1 week ago",
                "Benefits": ["Health Insurance", "Stock Options", "Flexible Hours"],
                "Qualifications": ["PhD in Statistics or related field", "Python", "Machine Learning"],
                "Job Description": "Join our data science team to solve complex business problems...",
                "Apply Link": "https://example.com/apply-ds",
                "Structured Content": {
                    "Requirements": ["Statistics background", "Programming skills", "Communication skills"],
                    "Responsibilities": ["Build models", "Analyze data", "Present findings"]
                }
            },
            # Add more sample jobs here
        ]
    }



def perform_cv_analysis(text):
    st.markdown('<div class="section-title">CV/Resume Analysis</div>', unsafe_allow_html=True)
    
    # Extract main CV components
    cv_data = {
        "Name": extract_name(text),
        "Email": extract_email(text),
        "Phone": extract_phone(text),
        "Location": extract_location(text),
        "Profile Summary": extract_profile_summary(text),
        "Experience": extract_experience(text),
        "Skills": extract_skills(text),
        "Qualifications": extract_qualifications(text),
        "Employment History": extract_employment_history(text),
        "Projects": extract_projects(text),
        "Languages": extract_languages(text),
        "Certifications": extract_certifications(text),
        "Social Links": extract_social_links(text)
    }

    # Create columns layout
    col1, col2 = st.columns(2)

    with col1:
        # Display Personal Information
        st.markdown('<div class="info-box"><b>Personal Information</b></div>', unsafe_allow_html=True)
        st.markdown(f"**Name:** {cv_data['Name']}")
        st.markdown(f"**Email:** {cv_data['Email']}")
        st.markdown(f"**Phone:** {cv_data['Phone']}")
        st.markdown(f"**Location:** {cv_data['Location']}")

        # Work history section
        st.markdown('<div class="info-box"><b>Employment History</b></div>', unsafe_allow_html=True)
        if isinstance(cv_data['Employment History'], list):
            for job in cv_data['Employment History']:
                st.markdown(f"- {job}")
        else:
            st.markdown(cv_data['Employment History'])

        # Display Languages
        st.markdown('<div class="info-box"><b>Languages</b></div>', unsafe_allow_html=True)
        if isinstance(cv_data['Languages'], list):
            for lang in cv_data['Languages']:
                st.markdown(f"- {lang}")

    with col2:
        # Profile summary
        st.markdown('<div class="info-box"><b>Profile Summary</b></div>', unsafe_allow_html=True)
        st.markdown(cv_data['Profile Summary'])

        # Experience overview
        st.markdown('<div class="info-box"><b>Experience</b></div>', unsafe_allow_html=True)
        st.markdown(f"{cv_data['Experience']}")

        # Skills visualization
        if isinstance(cv_data['Skills'], list) and len(cv_data['Skills']) > 0:
            st.markdown('<div class="info-box"><b>Key Skills</b></div>', unsafe_allow_html=True)
            skills_df = pd.DataFrame({
                "Skill": cv_data['Skills'],
                "Count": [1] * len(cv_data['Skills'])  # Just for visualization
            })
            fig = px.bar(skills_df, x="Skill", y="Count", title="Skills")
            st.plotly_chart(fig)

    # Create tabs for additional sections
    tabs = st.tabs(["Projects", "Qualifications", "Certifications"])

    with tabs[0]:
        st.markdown('<div class="info-box"><b>Projects</b></div>', unsafe_allow_html=True)
        if isinstance(cv_data['Projects'], list):
            for project in cv_data['Projects']:
                st.markdown(f"- {project}")

    with tabs[1]:
        st.markdown('<div class="info-box"><b>Qualifications & Education</b></div>', unsafe_allow_html=True)
        if isinstance(cv_data['Qualifications'], list):
            for qual in cv_data['Qualifications']:
                st.markdown(f"- {qual}")

    with tabs[2]:
        st.markdown('<div class="info-box"><b>Certifications</b></div>', unsafe_allow_html=True)
        if isinstance(cv_data['Certifications'], list):
            for cert in cv_data['Certifications']:
                st.markdown(f"- {cert}")

    return cv_data

if __name__ == "__main__":
   
    main()



